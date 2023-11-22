from ast import excepthandler
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import glob
import os
from pathlib import Path
#from pdf2docx import Converter

import subprocess, os, platform
import shutil


# def doc2docx(filepath, filename):
#     old_dir = os.path.dirname(filepath)

#     if (dest_dir == old_dir):
#         cmd = f'/Volumes/LibreOffice/LibreOffice.app/Contents/MacOS/soffice ' \
#             f'--headless --convert-to docx "{filename}"' 
#     else:    
#         cmd = f'/Volumes/LibreOffice/LibreOffice.app/Contents/MacOS/soffice ' \
#             f'--headless --convert-to docx --outdir "{dest_dir}" "{filename}"'
         
#     subprocess.run(cmd, cwd=old_dir, shell=True)
    

def move_file(filepath):
    filename = os.path.basename(filepath)

    if (filename[-3:] == "pdf"):
        return False
        # cv = Converter(filepath)
        # new_filepath = os.path.join(dest_dir, filename[:-3] + "docx")
        # cv.convert(new_filepath)
        # cv.close()
        # return new_filepath
    
    if (filename[-3:] == "doc"):
        doc2docx(filepath, filename)
        return os.path.join(dest_dir, filename + "x")
    
    if (dest_dir == os.path.dirname(filepath)):
        return filepath

    new_filepath = os.path.join(dest_dir, filename)
   
    if os.path.isfile(filepath):
        shutil.copy(filepath, new_filepath)

    return new_filepath



def reformat(filepath):
    print("Formatting " + os.path.basename(filepath) + "..........")

    resume = Document(filepath)
    header = resume.sections[0].header
    header.is_linked_to_previous = True
    paragraph = header.paragraphs[0]
    p = paragraph._p
    p.getparent().remove(p)
    htable = header.add_table(1, 2, Inches(7))

    htab_cells = htable.rows[0].cells

    ht_left = htab_cells[0].paragraphs[0]

    #Get Header_DB.py path with Path(__file__, '..').resolve()
    script_dir = os.path.dirname(__file__)
    logo_run = ht_left.add_run()
    logo_location = os.path.join(script_dir, "JRSS1.png")
    logo_run.add_picture(logo_location, width = Inches(1)) 
    logo_run.alignment = WD_ALIGN_PARAGRAPH.LEFT

    ht_right = htab_cells[1].paragraphs[0]

    slogan_run = ht_right.add_run()
    slogan_location = os.path.join(script_dir, "JRSS_FBanner.png")
    slogan_run.add_picture(slogan_location, width = Inches(4.5))
    slogan_run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    

    resume.save(filepath)



def run_db_directory():
    global dest_dir

    print("Welcome to the automated resume formatter for JRSS")
    #Add init so people can just click and run the file
    #Change so path is universal for whoever is running
    
    root = ""
    
    for subdir in os.listdir(root):
        subdir_path = os.path.join(root, subdir)

        if os.path.basename(subdir_path) == ".DS_Store":
            continue
        
        dest_dir = os.path.join(subdir_path, "Formatted")
        try:
            os.mkdir(dest_dir)
            print("Directory " + dest_dir + " made")

        except FileExistsError:
            pass

        for filename in os.listdir(subdir_path):
            filepath = os.path.join(subdir_path, filename)
            #check_file = os.path.join(dest_dir, filename)

            if os.path.isdir(filepath):
                continue

            if os.path.basename(filepath) == ".DS_Store":
                continue

            # if (os.path.isfile(check_file)):
            #     continue

            print("Filepath: " + filename)
            new_filepath = move_file(filepath)
            if new_filepath == False:
                continue
            print("File moved to " + new_filepath)
            
            reformat(new_filepath)
            print("Done")

    print("---------------------")
    print("All resumes formatted.")


def run_db_folder():
    global dest_dir

    print("Welcome to the automated resume formatter for JRSS")
    #Add init so people can just click and run the file
    #Change so path is universal for whoever is running
    root = input("Path to folder: ")
    dest_in = input("Path to dest: ")

    if (dest_in == ""):
        dest = None

    if (dest == None):
        dest_dir = os.path.join(root, "Formatted")
        try:
            os.mkdir(dest_dir)
            print("Directory " + dest_dir + " made")

        except FileExistsError:
            pass

    else:
        dest_dir = dest

    for filename in os.listdir(root):
        filepath = os.path.join(root, filename)

        if os.path.isdir(filepath):
                continue
    
        if os.path.basename(filepath) == ".DS_Store":
            continue

        print("Filepath: " + filename)
        new_filepath = move_file(filepath)
        if new_filepath == False:
            continue
        print("File moved to " + new_filepath)
        
        reformat(new_filepath)
        print("Done")

    print("---------------------")
    print("All resumes formatted.")

def run_db_file():
    global dest_dir
    global same_dir

    filepath = input("Path to file: ")
    dest_in = input("Path to dest: ")

    if (dest_in == ""):
        dest = None

    if (dest == None):
        dest_dir = os.path.dirname(filepath)
        same_dir = True

    else:
        dest_dir = dest

    new_filepath = move_file(filepath)
    if (not same_dir):
        print("File moved to " + new_filepath)
        
    reformat(new_filepath)
    print("Done")

def run_file():
    filepath = input("Path to file: ")
    reformat(filepath)
    print("File reformatted")
    

# run_db_folder(root, None)

# parent = os.path.dirname(filepath)
# dest = os.path.join(parent, "Formatted")

run_file()