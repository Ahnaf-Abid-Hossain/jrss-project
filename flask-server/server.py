from flask import Flask, request, jsonify, send_file
import os
from flask_cors import CORS  # Import CORS
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
from PIL import Image
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
import zipfile
import time
import threading
from datetime import datetime

app = Flask(__name__)

# Enable CORS for all routes
CORS(app)

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def modify_docx(file_path):
    # Load the document
    doc = Document(file_path)
    print("4")
    # Move header and footer content to the top of the body
    for section in doc.sections:
        header = section.header
        footer = section.footer
        header.is_linked_to_previous = True
        section.different_first_page_header_footer = False

        # move everything from header and footer to body
        # doc.paragraphs[0].insert_paragraph_before("--------------------------------\n")     
        # for paragraph in header.paragraphs:
        #     doc.paragraphs[0].insert_paragraph_before(paragraph.text)
        # doc.paragraphs[0].insert_paragraph_before("Header:")

        # doc.paragraphs[0].insert_paragraph_before("--------------------------------\n")
        # for paragraph in footer.paragraphs:
        #     doc.paragraphs[0].insert_paragraph_before(paragraph.text)
        # doc.paragraphs[0].insert_paragraph_before("Footer:")

        # Clear header and footer
        print("5")
        for paragraph in header.paragraphs:
            paragraph.clear()
        for paragraph in footer.paragraphs:
            paragraph.clear()
        print(len(header.paragraphs))
        print(len(footer.paragraphs))



        # Clear the existing content of the header
        # for paragraph in header.paragraphs:
        #     for run in paragraph.runs:
        #         run.clear()
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_paragraph.paragraph_format.left_indent = Inches(0)
        header_paragraph.paragraph_format.right_indent = Inches(0)

        header_paragraph.paragraph_format.space_before = Inches(0)
        header_paragraph.paragraph_format.space_after = Inches(0)   
        
        print("6")
        # set the indent of the table in the header (fixed the indentation bug)
        htable = header.add_table(1, 2, width=Inches(7))
        for row in htable.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.left_indent = Inches(0)



        # # Set individual column widths
        htable.columns[0].width = Inches(1)
        htable.columns[1].width = Inches(4.5)

        # Left cell (Logo)
        ht_left = htable.cell(0, 0).paragraphs[0]
        ht_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = ht_left.add_run()
        logo_run.add_picture("JRSS1.png", width=Inches(1))

        # Right cell (Slogan)
        ht_right = htable.cell(0, 1).paragraphs[0]
        ht_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        slogan_run = ht_right.add_run()
        slogan_run.add_picture("JRSS_FBanner.png", width=Inches(4.5))
        print("6")

    # Regex patterns for phone numbers and emails
    phone_pattern = r'^\s*(?:\+?(\d{1,3}))?[-. ()]*(\d{3})[-. )(]*(\d{3})[-. ]*(\d{4})(?: *x(\d+))?\s*$'  # Adjust this pattern as needed
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'  # Adjust as needed

    # Delete phone numbers and emails from the body
    print("7")
    for paragraph in doc.paragraphs:
        paragraph.text = re.sub(phone_pattern, '', paragraph.text)
        paragraph.text = re.sub(email_pattern, '', paragraph.text)
    print("8")
    # Save the modified document
    # doc.save(f"./modified_{file_path}")
    print(file_path)
    fileName = file_path.split("\\")[1]
    print(fileName)
    fileName = fileName.split(".")[0]
    print("9")
    doc.save(f"./uploads/{fileName}_modified.docx")
    print("here")
    # delete the original file
    # for cases, where file name ends in modified.docx it will be included in zip file
    pathToDelete = f"./uploads/{fileName}.docx"
    if os.path.exists(pathToDelete):
        os.remove(pathToDelete)

def schedule_file_deletion(file_path, delay=20):
    time.sleep(delay)
    os.remove(file_path)
    


@app.route('/upload', methods=['POST'])
def upload_file():
    # check if the post request has the file part
    if len(request.files) == 0:
        return jsonify({'error': 'No selected file'})
    
    fileNames = []
    current_unix_time = int(time.time())
    current_timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

    # save all the files
    for reqFiles in request.files:
        print("dobby")
        print(reqFiles)
        file = request.files[reqFiles]
        filename = os.path.join(app.config['UPLOAD_FOLDER'], f"{current_timestamp}{chr(300)}{str(reqFiles)}{chr(300)}{file.filename}")
        print(filename)
        file.save(filename)
        fileNames.append(filename)
        
    # modify all the documents
    try:
        for name in fileNames:
            modify_docx(name)
    except:
        return jsonify({'error': 'Error modifying document'})
    
        


    directory = "./uploads"

    # create the zip file with all the modified documents
    with zipfile.ZipFile(f'{directory}/{current_unix_time}_modified_documents.zip', 'w') as zipObj:
        for folderName, subfolders, filenames in os.walk(directory):
            for filename in filenames:
                # if filename.startswith(f"{current_unix_time}"):
                if filename.endswith("modified.docx"):
                    relative_path = os.path.relpath(os.path.join(folderName, filename), directory)
                    splitted_name = relative_path.split(chr(300))
                    zipObj.write(os.path.join(folderName, filename), splitted_name[-2] + splitted_name[-1])
                
                    pathToDelete = os.path.join(folderName, filename)
    
                    # Check if the file exists to prevent error
                    if os.path.exists(pathToDelete):
                        print("Deleting file: ", pathToDelete)
                        os.remove(pathToDelete)


    # Schedule the file for deletion after 20 seconds
    threading.Thread(target=schedule_file_deletion, args=(f'{directory}/{current_unix_time}_modified_documents.zip', 5)).start()
    
    # create some sort of background job to delete the files after a certain amount of time
    # or delete the file after a delay?
    return send_file(f'./uploads/{current_unix_time}_modified_documents.zip', as_attachment=True)



if __name__ == '__main__':
    app.run(port=8000, debug=True)
    
    